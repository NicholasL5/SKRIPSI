import json
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

"""
This script is used for converting dataset.json file into docx file, with cleaned latex formatting.
"""

def sanitize_for_xml(text):
    """Remove or replace characters that are not compatible with XML."""
    if not text:
        return text
        
    result = ""
    for char in text:
        if char in ['\t', '\n', '\r', ' ']:
            result += char
        elif ord(char) < 32:
            continue
        else:
            result += char
    
    return result

def clean_latex(text):
    """Convert LaTeX expressions to regular text."""
    # Complex regex replacements (with capture groups) - do these first
    regex_replacements = [
        (r'\\left\(', r'('),
        (r'\\right\)', r')'),
        (r'\\left\[', r'['),
        (r'\\right\]', r']'),
        (r'\\left\\{', r'{'),
        (r'\\right\\}', r'}'),
        (r'\\text\{([^}]+)\}', r'\1'),
        (r'\\sqrt\{([^}]+)\}', r'√(\1)'),
        # Handle general fractions with any numerator/denominator
        (r'\\frac\{([^}]+)\}\{([^}]+)\}', r'(\1)/(\2)'),
        # Special case for tfrac
        (r'\\tfrac\{([^}]+)\}\{([^}]+)\}', r'(\1)/(\2)'),
        # Handle overset with frown (arc notation)
        (r'\\overset{\\frown}\{([^}]+)\}', r'⌢\1'),
        (r'\s*\^3', '³'),
        (r'\s*\^2', '²'),
    ]
    
    for pattern, replacement in regex_replacements:
        text = re.sub(pattern, replacement, text)
    
    # Simple regex
    simple_replacements = {
        r'\times': '×',
        r'\cdot': '·',      
        r'\quad': '    ',  
        r'\Rightarrow': '⇒',
        r'\pi': 'π',      
        r'\frown': '⌢', 
        r'\theta': 'θ',
        r'\circ': '°',
        r'\alpha': 'α',
        r'\rightarrow': '→',
        r'\implies': '⇒',
        r'\angle': '∠',
        r'\delta': '∆',
        r'\Delta': '∆',
        r'\,': ' ',
        r'\tfrac{1}{2}': '½',
        r'\Bigl': '',
        r'\Bigr': '',
        r'\\': '\n',
        r'\n': '\n',
        r'\[': '',
        r'\]': '',
        r'\(': '',
        r'\)': '',
        r'\sum': 'Σ',
        r'\neq': '≠',
        r'\approx': '≈',
        r'\equiv': '≡',     
        r'\sim': '∼',       
        r'\le': '≤',         
        r'\leq': '≤',  
        r'\ge': '≥',         
        r'\geq': '≥',
    }
    
    # Apply simple string replacements
    for pattern, replacement in simple_replacements.items():
        text = text.replace(pattern, replacement)
    
    return text

def format_conversation(conversations, output_file):
    doc = Document()
    
    title = doc.add_heading('Matematika Conversation', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    for i, conversation in enumerate(conversations):
        doc.add_heading(f'Conversation {i+1}', 1)
        
        for message in conversation["data"]:
            role = message["role"]
            content = clean_latex(message["content"])
            
            p = doc.add_paragraph()
            if role == "user":
                p.add_run("User: ").bold = True
            else:
                p.add_run("Assistant: ").bold = True
            
            # Add content
            p.add_run(sanitize_for_xml(content))
            
        if i < len(conversations) - 1:
            doc.add_page_break()
    
    doc.save(output_file)
    print(f"Conversation saved to {output_file}")

def main():
    input_file = "dataset - bangun-ruang.json"
    output_file = "math_conversations.docx"
    
    try:
        with open(input_file, 'r', encoding='utf-8') as f:
            conversations = json.load(f)
            
            if not isinstance(conversations, list):
                conversations = [conversations]
        
        format_conversation(conversations, output_file)
        
    except Exception as e:
        print(f"Error: {e}")

if __name__ == "__main__":
    main()